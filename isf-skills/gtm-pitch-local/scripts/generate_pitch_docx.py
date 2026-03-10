"""
Reusable GTM Pitch Word Document Generator
==========================================
Single entry point: generate_pitch_docx(data, output_path)

`data` is a plain dict with the following structure. All keys are optional
except company_name and industry — the renderer skips any section whose
key is missing or empty.

Schema
------
{
  # ── Title page ──────────────────────────────────────────────
  "company_name": str,
  "industry": str,
  "persona_title": str,          # e.g. "Chief Marketing Officer (CMO)"
  "sfdc_id": str,
  "account_line": str,           # e.g. "Tier 1 | Majors | G2K #1191 | AE: … | Existing Customer"

  # ── Optional alert banner (red) ─────────────────────────────
  "alert": str | None,

  # ── TIER 1 ──────────────────────────────────────────────────
  "why_snowflake": [              # list of (title, description) tuples
      ("Reason 1 Title", "Full paragraph …"),
      …
  ],

  "why_coco": [                   # list of (pillar_label, description)
      ("\U0001F680 Faster Innovation", "…"),
      ("\U0001F9E0 Deep Awareness", "…"),
      ("\U0001F3DB\uFE0F Enterprise-Ready", "…"),
  ],
  "coco_quote": ("quote text", "attribution", "source_url"),   # single featured quote (url optional)

  "top_solutions": [              # exactly 3 for TIER 1 table
      {"name": str, "id": str, "value": str, "products": str},
      …
  ],

  "pain_points": [                # list of (title, url, description)
      ("Pain title", "https://…/pain/PAIN-xxx", "Context for this account"),
      …
  ],

  "competitive": {
      "competitors": str,
      "approach": str,
      "edge": str,
  },

  "proof_point": {
      "title": str,
      "url": str,
      "summary": str,
  },

  "persona": {                     # key-value pairs rendered in order
      "fields": [("Label", "Value"), …],
      "isf_link": str | None,
  },

  # ── TIER 2 ──────────────────────────────────────────────────
  "industry_context": {
      "trends_label": str,         # e.g. "Retail Market Trends:"
      "trends": [str, …],
      "company_context_label": str, # e.g. "DoorDash-Specific Context:"
      "company_context": [str, …],
      "snowflake_pov": str,
      "industry_id": str,          # for ISF link
      "industry_name": str,
  },

  "all_solutions": [               # full list for TIER 2 table
      {"name": str, "id": str, "value": str, "products": str},
      …
  ],

  "competitive_deep_dive": {
      "tech_stack": str,
      "signals": [str, …],
      "strategies": [("Label: ", "Description"), …],
  },

  "stories": [                     # list of dicts
      {"title": str, "id": str, "customer": str, "outcomes": str,
       "public_url": str | None},
      …
  ],

  "coco_detail": {
      "pillars": [                 # list of (pillar_title, [bullet, …])
          ("Pillar 1 — \U0001F680 Faster Innovation:", ["…", "…"]),
          …
      ],
      "benchmarks": str,
      "quotes": [("quote", "author", "relevance_note", "source_url"), …],
  },

  "coco_prompts": [                # list of (title, prompt, what_it_builds)
      ("Pitch Idea 1: …", '"CoCo prompt…"', "Description of output"),
      …
  ],

  "recommended_personas": [        # list of (name, description)
      ("VP of Marketing", "…"),
      …
  ],

  "coco_skills_appendix": [         # list of dicts for appendix table
      {"name": str, "type": str, "pillar": str, "description": str, "relevance": str},
      …
  ],

  "sources": str,                  # footer source attribution
}
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

ISF_BASE = 'https://kcxcbaixd-sfcogsops-snowhouse-aws-us-west-2.snowflakecomputing.app'

NAVY = RGBColor(0x1B, 0x3A, 0x6B)
ACCENT = RGBColor(0x29, 0xB5, 0xE8)
DARK = RGBColor(0x33, 0x33, 0x33)
GREY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GREY = RGBColor(0x99, 0x99, 0x99)
RED = RGBColor(0xCC, 0x00, 0x00)
GREEN = RGBColor(0x0A, 0x66, 0x0A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BORDER_GREY = RGBColor(0xCC, 0xCC, 0xCC)

COCO_LINKS = [
    ('Product Page', 'https://www.snowflake.com/en/product/features/cortex-code/'),
    ('Blog', 'https://www.snowflake.com/en/blog/cortex-code-cli-expands-support/'),
    ('Compass', 'https://snowflake.seismic.com/Link/Content/DCq6VH9PcX6h3GWPVGhTgVPgXqHd'),
    ('Docs', 'https://docs.snowflake.com/en/user-guide/cortex-code/cortex-code-cli'),
    ('Try Free', 'https://signup.snowflake.com/cortex-code'),
]

COCO_LINKS_EXTENDED = COCO_LINKS + [
    ('GA Press Release', 'https://www.snowflake.com/en/news/press-releases/snowflake-unveils-cortex-code-an-ai-coding-agent-that-drastically-increases-productivity-by-understanding-your-enterprise-data-context/'),
]


def _add_run(para, text, bold=False, italic=False, color=None, size=None, font_name=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    if font_name:
        run.font.name = font_name
    return run


def _add_hyperlink(paragraph, display_text, url, bold=False, size=None, font_name='Calibri'):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '2980B9')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    if size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size * 2))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(size * 2))
        rPr.append(szCs)
    new_run.append(rPr)
    new_run.text = display_text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_hyperlink_to_cell(cell, display_text, url, bold=False, size=None):
    para = cell.paragraphs[0]
    if para.text:
        para = cell.add_paragraph()
    _add_hyperlink(para, display_text, url, bold=bold, size=size)


def _set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def _style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = 'Calibri'
            if row_idx == 0:
                _set_cell_shading(cell, '1B3A6B')
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = WHITE
                        run.bold = True


def _heading(doc, text, level=1, color=NAVY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def _coco_resource_links(doc, extended=False):
    p = doc.add_paragraph()
    _add_run(p, 'Resources: ', bold=True, color=NAVY)
    links = COCO_LINKS_EXTENDED if extended else COCO_LINKS
    for idx, (name, url) in enumerate(links):
        _add_hyperlink(p, name, url)
        if idx < len(links) - 1:
            _add_run(p, ' | ')
    return p


def _solution_url(sol_id):
    return f'{ISF_BASE}/solution/{sol_id}'


def _story_url(story_id):
    return f'{ISF_BASE}/story/{story_id}'


def _solutions_table(doc, solutions):
    headers = ['Solution', 'Value Positioning', 'Products', 'Links']
    t = doc.add_table(rows=1 + len(solutions), cols=4)
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        t.cell(0, i).text = h
    for row_idx, sol in enumerate(solutions, 1):
        p = t.cell(row_idx, 0).paragraphs[0]
        _add_hyperlink(p, sol['name'], _solution_url(sol['id']), bold=True)
        t.cell(row_idx, 1).text = sol['value']
        t.cell(row_idx, 2).text = sol['products']
        _add_hyperlink_to_cell(t.cell(row_idx, 3), 'View full brief', _solution_url(sol['id']))
    _style_table(t)


def _stories_table(doc, stories):
    headers = ['Story', 'Customer', 'Key Outcomes', 'Links']
    t = doc.add_table(rows=1 + len(stories), cols=4)
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        t.cell(0, i).text = h
    for row_idx, st in enumerate(stories, 1):
        p = t.cell(row_idx, 0).paragraphs[0]
        _add_hyperlink(p, st['title'], _story_url(st['id']), bold=True)
        t.cell(row_idx, 1).text = st['customer']
        t.cell(row_idx, 2).text = st['outcomes']
        cell3 = t.cell(row_idx, 3)
        _add_hyperlink_to_cell(cell3, 'ISF', _story_url(st['id']))
        pub = st.get('public_url')
        if pub:
            _add_run(cell3.paragraphs[0], ' | ', size=Pt(9))
            _add_hyperlink(cell3.paragraphs[0], 'Public', pub, size=9)
    _style_table(t)


def generate_pitch_docx(data: dict, output_path: str) -> str:
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = DARK

    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.color.rgb = NAVY
        hs.font.name = 'Calibri'

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    company = data['company_name']
    industry = data['industry']

    # ── TITLE PAGE ─────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(t, '\U0001F3AF', size=Pt(36))

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(t2, 'GTM Pitch', bold=True, color=NAVY, size=Pt(28))

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(t3, company, bold=True, color=ACCENT, size=Pt(24))

    t4 = doc.add_paragraph()
    t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(t4, industry, color=GREY, size=Pt(14))

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if data.get('persona_title'):
        _add_run(meta, 'Persona: ', bold=True, color=NAVY)
        _add_run(meta, data['persona_title'] + '\n')
    _add_run(meta, 'Generated: ', bold=True, color=NAVY)
    _add_run(meta, datetime.date.today().strftime("%B %d, %Y") + '\n')
    if data.get('sfdc_id'):
        _add_run(meta, 'SFDC ID: ', bold=True, color=NAVY)
        _add_run(meta, data['sfdc_id'] + '\n')
    if data.get('account_line'):
        _add_run(meta, 'Account: ', bold=True, color=NAVY)
        _add_run(meta, data['account_line'])

    doc.add_page_break()

    # ── ALERT BANNER ───────────────────────────────────────────
    if data.get('alert'):
        alert = doc.add_paragraph()
        _add_run(alert, '\u26A0\uFE0F CRITICAL CONTEXT: ', bold=True, color=RED)
        _add_run(alert, data['alert'])

    # ── WHY SNOWFLAKE ──────────────────────────────────────────
    if data.get('why_snowflake'):
        _heading(doc, f'Why Snowflake for {company}')
        for i, (title, desc) in enumerate(data['why_snowflake'], 1):
            p = doc.add_paragraph()
            _add_run(p, f'{i}. {title}', bold=True, color=NAVY)
            _add_run(p, f' \u2014 {desc}')

    # ── WHY COCO CLI ───────────────────────────────────────────
    if data.get('why_coco'):
        _heading(doc, f'Why CoCo CLI for {company}')
        for i, (pillar, desc) in enumerate(data['why_coco'], 1):
            p = doc.add_paragraph()
            _add_run(p, f'{i}. {pillar}', bold=True, color=NAVY)
            _add_run(p, f' \u2014 {desc}')

        if data.get('coco_quote'):
            cq = data['coco_quote']
            q_text = cq[0]
            q_attr = cq[1]
            q_source = cq[2] if len(cq) > 2 else None
            qp = doc.add_paragraph()
            qp.paragraph_format.left_indent = Cm(1)
            _add_run(qp, q_text, italic=True)
            _add_run(qp, '\n\u2014 ', bold=True)
            _add_run(qp, q_attr, bold=True)
            if q_source:
                _add_run(qp, ' | ')
                _add_hyperlink(qp, 'Source', q_source, size=9)

        _coco_resource_links(doc)

    # ── TOP 3 SOLUTIONS ────────────────────────────────────────
    if data.get('top_solutions'):
        _heading(doc, 'Top 3 Solutions')
        _solutions_table(doc, data['top_solutions'][:3])

    # ── PAIN POINTS ────────────────────────────────────────────
    if data.get('pain_points'):
        _heading(doc, 'Top 3 Pain Points to Lead With')
        for idx, (title, url, desc) in enumerate(data['pain_points'][:3], 1):
            p = doc.add_paragraph()
            _add_run(p, f'{idx}. ', bold=True)
            if url:
                _add_hyperlink(p, title, url, bold=True)
            else:
                _add_run(p, title, bold=True)
            _add_run(p, f' \u2014 {desc}')

    # ── COMPETITIVE ────────────────────────────────────────────
    if data.get('competitive'):
        comp = data['competitive']
        _heading(doc, 'Competitive Positioning')
        for label, key in [('Primary Competitors: ', 'competitors'),
                           ('Approach: ', 'approach'),
                           ('Our Edge: ', 'edge')]:
            if comp.get(key):
                p = doc.add_paragraph()
                _add_run(p, label, bold=True, color=NAVY)
                _add_run(p, comp[key])

    # ── PROOF POINT ────────────────────────────────────────────
    if data.get('proof_point'):
        pp = data['proof_point']
        _heading(doc, 'Proof Point')
        p = doc.add_paragraph()
        _add_hyperlink(p, pp['title'], pp['url'], bold=True)
        _add_run(p, f' \u2014 {pp["summary"]}')
        lp = doc.add_paragraph()
        _add_run(lp, 'Links: ', bold=True)
        _add_hyperlink(lp, 'ISF Detail', pp['url'])
        if pp.get('public_url'):
            _add_run(lp, ' | ')
            _add_hyperlink(lp, 'Public Case Study', pp['public_url'])

    # ── PERSONA ────────────────────────────────────────────────
    if data.get('persona'):
        persona = data['persona']
        title = data.get('persona_title', 'Persona')
        _heading(doc, f'Persona: {title}')
        if persona.get('isf_link'):
            lp = doc.add_paragraph()
            _add_hyperlink(lp, 'View in ISF', persona['isf_link'], size=9)
        for label, value in persona.get('fields', []):
            p = doc.add_paragraph()
            _add_run(p, f'{label}: ', bold=True, color=NAVY)
            _add_run(p, value)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # TIER 2: DETAILED PITCH
    # ════════════════════════════════════════════════════════════
    _heading(doc, 'Detailed Pitch')

    # ── INDUSTRY CONTEXT ───────────────────────────────────────
    if data.get('industry_context'):
        ic = data['industry_context']
        _heading(doc, 'Industry Context', level=2)

        if ic.get('trends'):
            tp = doc.add_paragraph()
            _add_run(tp, ic.get('trends_label', 'Market Trends:'), bold=True, color=NAVY)
            for t in ic['trends']:
                doc.add_paragraph(t, style='List Bullet')

        if ic.get('company_context'):
            cp = doc.add_paragraph()
            _add_run(cp, ic.get('company_context_label', f'{company}-Specific Context:'), bold=True, color=NAVY)
            for c in ic['company_context']:
                doc.add_paragraph(c, style='List Bullet')

        if ic.get('snowflake_pov'):
            p = doc.add_paragraph()
            _add_run(p, 'Snowflake POV: ', bold=True, color=NAVY)
            _add_run(p, ic['snowflake_pov'])

        if ic.get('industry_id'):
            lp = doc.add_paragraph()
            name = ic.get('industry_name', 'Industry')
            _add_hyperlink(lp, f'View {name} in ISF', f'{ISF_BASE}/industry/{ic["industry_id"]}')

    # ── ALL SOLUTIONS ──────────────────────────────────────────
    if data.get('all_solutions'):
        _heading(doc, 'All Recommended Solutions', level=2)
        _solutions_table(doc, data['all_solutions'])

    # ── COMPETITIVE DEEP DIVE ──────────────────────────────────
    if data.get('competitive_deep_dive'):
        cd = data['competitive_deep_dive']
        _heading(doc, 'Competitive Deep Dive', level=2)

        if cd.get('tech_stack'):
            p = doc.add_paragraph()
            _add_run(p, 'Current Tech Stack: ', bold=True, color=NAVY)
            _add_run(p, cd['tech_stack'])

        if cd.get('signals'):
            p = doc.add_paragraph()
            _add_run(p, 'Competitor Signals:', bold=True, color=NAVY)
            for s in cd['signals']:
                doc.add_paragraph(s, style='List Bullet')

        if cd.get('strategies'):
            p = doc.add_paragraph()
            _add_run(p, 'Displacement Strategy:', bold=True, color=NAVY)
            for label, desc in cd['strategies']:
                sp = doc.add_paragraph()
                _add_run(sp, label, bold=True, color=NAVY)
                _add_run(sp, desc)

    # ── CUSTOMER STORIES ───────────────────────────────────────
    if data.get('stories'):
        _heading(doc, 'Customer Stories', level=2)
        _stories_table(doc, data['stories'])

    # ── COCO DETAIL ────────────────────────────────────────────
    if data.get('coco_detail'):
        cd = data['coco_detail']
        _heading(doc, 'CoCo CLI Pitch Detail', level=2)

        for pillar_title, bullets in cd.get('pillars', []):
            p = doc.add_paragraph()
            _add_run(p, pillar_title, bold=True, color=NAVY)
            for b in bullets:
                doc.add_paragraph(b, style='List Bullet')

        if cd.get('benchmarks'):
            p = doc.add_paragraph()
            _add_run(p, 'Benchmarks: ', bold=True, color=NAVY)
            _add_run(p, cd['benchmarks'])

        if cd.get('quotes'):
            _heading(doc, 'Customer Quotes', level=3)
            for q_item in cd['quotes']:
                q_text = q_item[0]
                q_author = q_item[1]
                q_note = q_item[2] if len(q_item) > 2 else ''
                q_source = q_item[3] if len(q_item) > 3 else None
                qp = doc.add_paragraph()
                qp.paragraph_format.left_indent = Cm(1)
                _add_run(qp, q_text, italic=True)
                _add_run(qp, '\n\u2014 ', bold=True)
                _add_run(qp, q_author, bold=True)
                if q_note:
                    _add_run(qp, f' ({q_note})', italic=True, color=GREY)
                if q_source:
                    _add_run(qp, ' | ')
                    _add_hyperlink(qp, 'Source', q_source, size=9)

        _coco_resource_links(doc, extended=True)

    # ── COCO PROMPTS ───────────────────────────────────────────
    if data.get('coco_prompts'):
        _heading(doc, 'Try It in CoCo CLI \u2014 Top 3 Pitch Ideas in Action', level=2)
        doc.add_paragraph('Paste these prompts directly into CoCo CLI to bring each pitch idea to life:')

        for title, prompt, desc in data['coco_prompts']:
            pp = doc.add_paragraph()
            _add_run(pp, title, bold=True, color=NAVY)

            pr = doc.add_paragraph()
            pr.paragraph_format.left_indent = Cm(1)
            _add_run(pr, prompt, font_name='Consolas', color=GREEN, size=Pt(9.5))

            dp = doc.add_paragraph()
            dp.paragraph_format.left_indent = Cm(1)
            _add_run(dp, 'What CoCo will build: ', bold=True, italic=True)
            _add_run(dp, desc, italic=True)

        ip = doc.add_paragraph()
        _add_run(ip, '\U0001F4A1 Getting started: ', bold=True, color=NAVY)
        _add_run(ip, 'Install CoCo CLI: ')
        _add_run(ip, 'curl -LsS https://ai.snowflake.com/static/cc-scripts/install.sh | sh',
                 font_name='Consolas', size=Pt(9))
        _add_run(ip, '\nOr sign up at ')
        _add_hyperlink(ip, 'https://signup.snowflake.com/cortex-code',
                       'https://signup.snowflake.com/cortex-code')
        _add_run(ip, ' ($40 free credits for 30 days)')

    # ── RECOMMENDED PERSONAS ───────────────────────────────────
    if data.get('recommended_personas'):
        _heading(doc, 'Recommended Personas to Engage', level=2)
        for name, desc in data['recommended_personas']:
            p = doc.add_paragraph()
            _add_run(p, f'{name}: ', bold=True, color=NAVY)
            _add_run(p, desc)

    # ── COCO SKILLS APPENDIX ─────────────────────────────────
    if data.get('coco_skills_appendix'):
        doc.add_page_break()
        _heading(doc, 'Appendix: CoCo Skills Reference', level=1)
        doc.add_paragraph('Native Cortex Code skills referenced in this pitch, plus related skills for this customer.')
        skills = data['coco_skills_appendix']
        tbl = doc.add_table(rows=1 + len(skills), cols=5)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(['Skill', 'Type', 'Pillar', 'Description', 'Relevance']):
            cell = tbl.rows[0].cells[i]
            cell.text = ''
            _add_run(cell.paragraphs[0], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(9))
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), '1B3A6B')
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(shading)
        for idx, sk in enumerate(skills):
            row = tbl.rows[idx + 1]
            vals = [sk.get('name',''), sk.get('type',''), sk.get('pillar',''), sk.get('description',''), sk.get('relevance','')]
            for ci, val in enumerate(vals):
                cell = row.cells[ci]
                cell.text = ''
                sz = Pt(8) if ci == 3 else Pt(9)
                _add_run(cell.paragraphs[0], val, size=sz)
        _style_table(tbl)

    # ── FOOTER ─────────────────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(footer, '\u2500' * 40 + '\n', color=BORDER_GREY)
    _add_run(footer,
             f'Generated: {datetime.datetime.now().strftime("%B %d, %Y at %I:%M %p")}\n',
             italic=True, color=LIGHT_GREY, size=Pt(8))
    if data.get('sources'):
        _add_run(footer, f'Sources: {data["sources"]}\n',
                 italic=True, color=LIGHT_GREY, size=Pt(8))
    _add_run(footer, 'Generated by Cortex Code GTM Pitch Skill',
             italic=True, color=LIGHT_GREY, size=Pt(8))

    doc.save(output_path)
    size = os.path.getsize(output_path)
    print(f'Saved: {output_path} ({size/1024:.1f} KB)')
    return output_path


if __name__ == '__main__':
    print('This module is imported by the GTM pitch skill. Usage:')
    print('  from generate_pitch_docx import generate_pitch_docx')
    print('  generate_pitch_docx(data_dict, "/path/to/output.docx")')
